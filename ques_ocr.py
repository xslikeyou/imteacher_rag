from pathlib import Path
from api import chat, set_client
from openai import RateLimitError
import streamlit as st
import os
from docx import Document
from io import BytesIO

def generate_docx(text: str) -> BytesIO:
    doc = Document()
    doc.add_heading("公式识别", level=1)
    for line in text.split('\n'):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def output_button(final_lesson):
    # 显示导出按钮区域
    st.markdown("### 📤 导出识别结果")
    if st.download_button("导出为 Word 文档", generate_docx(final_lesson), file_name="识别结果.docx"):
        st.success("Word 文件已准备好下载。")
# 创建一个文件上传器，允许用户上传文件
uploaded_file = st.file_uploader("请上传文件", type=["pdf", "doc", "png", "jpg", "jpeg", "bmp", "gif"])

# 检查是否有文件被上传
if uploaded_file:
    st.write("文件上传中...")
    # 构造完整的保存路径
    save_directory = './upload_file'
    os.makedirs(save_directory, exist_ok=True)
    save_path = os.path.join(save_directory, uploaded_file.name)
    # 保存文件到指定位置
    with open(save_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # 根据文件类型调用不同的API***********************************************************************
    if Path(uploaded_file.name).suffix.lower() in ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.pdf', '.doc']:
        try:
            print("获取图片")
            file_object = set_client.files.create(file=uploaded_file, purpose="file-extract")
            print("识别图片")
            file_content = set_client.files.content(file_id=file_object.id).text
        except RateLimitError as e:
            st.markdown(e)
            st.session_state.messages.append({"role": "assistant", "content": "error" + e})

    else:
        st.error("不支持的文件格式")
        file_content = ""

    messages = [
         {
            "role": "system",
            "content": file_content,
        },
        {"role": "user", 
         "content": r"""
                    你可以OCR识别下面图片中的文字，无论是手写体含还是打印体，
                    注意数学公式的输出格式：
                    \(行内公式\)
                    \[
                    行间公式
                    \]
                    你只输出图片中内容，不要输出提示语
                    """
         },
    ]

    # 然后调用 chat-completion, 获取 Kimi 的回答
    st.write("上传成功，题目识别...")
    completion = chat(history=messages)

    # 输出结果
    with st.chat_message("assistant", avatar='🤖'):
        st.write("识别latex代码")
        placeholder = st.empty()
        collected_messages = ""
        for chunk in completion:
            temp_messages = chunk.choices[0].delta.content
            if not temp_messages:
                continue
            collected_messages = collected_messages + temp_messages
            placeholder.code(collected_messages, language="latex")
        print(collected_messages)
        # 修改模型输出内容，进行latex格式矫正
        collected_messages = collected_messages
        collected_messages = collected_messages.replace(r"\(", "$").replace(r"\)", "$").replace(r"\[", "$$").replace(r"\]", "$$")
        placeholder.code(collected_messages, language="latex")
        # 将字符串转换为Markdown格式，并设置背景颜色
        st.write("latex渲染预览")
        md_with_background = """
        <div style="background-color: #f0f8ff; padding: 10px; border-radius: 5px;">
            {}
        </div>
        """.format(collected_messages)

        # 使用st.markdown渲染带有背景颜色的Markdown内容
        st.markdown(md_with_background, unsafe_allow_html=True)
        # 保存识别结果到 st.session_state
        output_button( collected_messages.replace("$", ""))


