
import streamlit as st
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from docx import Document
from io import BytesIO
from stream_answer import stream_ans
embedding_model="./all-MiniLM-L6-v2"
tplan_path='./tplan_db'

sys_prompt = """
       你擅长生成数学教案，教案结构：
       其中<3.教学目标4.教学重点7.教学过程> 是重点，其他可以省略
        1.内容分析:分析所教内容在数学学习中的地位
        2.学情分析：学生已经学习的知识，学生的接受能力等
        3.教学目标：一般三点，不能按三维目标写成知识与技能、过程与方法、情感态度与价值观，要每一点都尽量体现这三个维度
        4.教学重点：一般是2-3个重点
        5.教学难点：一般是一个难点即可
        6.教法学法：任务驱动法、师生互动法，等等
        7.教学过程：一定要写设计意图<这部分尽量详细>,要有具体的例题讲解，和学生习题练习，来源可以是课本，或者你自己命题目
        8.板书设计
        9.反思:教学完成情况
    """

def generate_docx(text: str) -> BytesIO:
    doc = Document()
    doc.add_heading("教案", level=1)
    for line in text.split('\n'):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def output_button(final_lesson):
    # 显示导出按钮区域
    st.markdown("### 📤 导出教案")
    if st.download_button("导出为 Word 文档", generate_docx(final_lesson), file_name="教案.docx"):
        st.success("Word 文件已准备好下载。")

def init_chat_history():
    if "tplan_db" not in st.session_state:
        st.session_state.tplan_db = []

def main():
    # 设置页面标题
    st.title("📘 教案生成助手")

    # 表单开始
    with st.form("lesson_plan_form"):
        subject = st.text_input("请输入课程主题：", placeholder="例如：集合的概念")
        version = st.selectbox("请选择教程版本：", ["人教版", "苏教版", "北师大版"])
        submit = st.form_submit_button("生成教案")
    init_chat_history()
    if submit:
        if not subject:
            st.warning("请先输入课程主题。")
        else:
            # 构造提示词
            prompt = "请根据以下信息生成一份详细的教案。\n\n课程主题：{}\n教程版本：{}\n\n教案内容：".format(subject,version)

            # 查询教案库、
            search_results_string=""
            if st.session_state.tplan_db:
                search_results=st.session_state.tplan_db.similarity_search(prompt, k=2)#result是document
            else:
                st.write("首次访问需加载教案库...请稍后")
                embeddings = HuggingFaceEmbeddings(model_name=embedding_model)
                vectorstore = FAISS.load_local(folder_path=tplan_path, embeddings=embeddings, allow_dangerous_deserialization=True)
                st.session_state.tplan_db = vectorstore
                st.write("教案库加载完成")
                search_results=st.session_state.tplan_db.similarity_search(prompt, k=2)
            for result in search_results:
                search_results_string += result.page_content + "\n"
            with st.expander("教案库查询结果"):
                st.write(search_results)
            history=[
                {"role": "system", "content": sys_prompt},
                {"role": "system", "content": search_results_string},
                {"role": "user", "content": prompt}
            ]
            lesson_plan=stream_ans(history)
            # 导出教案
            output_button(lesson_plan)

if __name__ == "__main__":
    main()